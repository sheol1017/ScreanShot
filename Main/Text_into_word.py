#!usr/bin/env/python3
#-*- coding: utf-8 -*-
__author__ = 'Michael Shen'

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
import os
#以上是docx库中需要用到的部分
import time
today1 = time.strftime("%Y-%m-%d",time.localtime())
today2 = time.strftime("%Y/%m/%d",time.localtime())
today = time.strftime("%Y{y}%m{m}%d{d}",time.localtime()).format(y='年',m='月',d='日')

def Text_word_init(address):
        if not os.path.exists('./result'):
                os.mkdir('./result')
        # with io.open('./result/test.doc', 'w') as fp:
        #         fp.write("%s" % (today))
        file=Document()
        # Document().save(address)
        file.add_paragraph('%s' %(today))
        file.save(address)

def Text_add_word(text,address):
        Orignal_file = Document(address)
        str=[]

        Orignal_file.add_paragraph("%s"%text)
        for para in Orignal_file.paragraphs:
            str.append(para.text)
        for s in str:
                print(s)

        # if not os.path.exists('./result'):
        #         os.mkdir('./result')
        Orignal_file.save(address)
        return