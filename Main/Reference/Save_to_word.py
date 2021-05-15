#!usr/bin/env/python3
#-*- coding: utf-8 -*-
__author__ = 'Michael Shen'
#pip install python-docx
from docx import Document
from docx.shared import Inches

string = '文字内容'
images = 'haha.jpg'  # 保存在本地的图片
doc = Document()  # doc对象
doc.add_paragraph(string)  # 添加文字
doc.add_picture(images, width=Inches(2))  # 添加图, 设置宽度
doc.save('word文档.docx')  # 保存路径